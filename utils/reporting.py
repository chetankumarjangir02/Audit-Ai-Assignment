from docx import Document
import pandas as pd
from datetime import datetime
import os

class ReportGenerator:
    def __init__(self, db):
        self.db = db

    def generate_audit_report(self, document_id=None, output_path=None):
        if not output_path:
            output_path = f"./outputs/reports/audit_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        doc = Document()
        doc.add_heading('Financial Audit Report', 0)
        doc.add_heading('Executive Summary', level=1)
        doc.add_paragraph(self._generate_executive_summary(document_id))
        doc.add_heading('Transaction Analysis', level=1)
        self._add_transaction_analysis(doc, document_id)
        doc.add_heading('Categories Breakdown', level=1)
        self._add_categories_analysis(doc, document_id)
        doc.save(output_path)
        return output_path

    def _generate_executive_summary(self, document_id=None):
        if document_id:
            labeled = list(self.db.labeled_data.find({"source_document_id": document_id}))
            total_transactions = len(labeled)
            total_docs = 1
        else:
            total_transactions = self.db.labeled_data.count_documents({})
            total_docs = self.db.documents.count_documents({})
        labeled = list(self.db.labeled_data.find()) if not document_id else list(self.db.labeled_data.find({"source_document_id": document_id}))
        df = pd.DataFrame(self._clean_numeric_data(labeled)) if labeled else pd.DataFrame()
        total_debit = df['debit'].sum() if 'debit' in df.columns and not df.empty else 0
        total_credit = df['credit'].sum() if 'credit' in df.columns and not df.empty else 0
        net_flow = total_credit - total_debit
        summary = f"""This audit report summarizes the financial transactions processed by the Audit Intelligence System.

System Overview:
- Total Documents Processed: {total_docs}
- Total Transactions Analyzed: {total_transactions}
- Total Debit Amount: ${total_debit:,.2f}
- Total Credit Amount: ${total_credit:,.2f}
- Net Cash Flow: ${net_flow:,.2f}
- Report Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
"""
        return summary

    def _add_transaction_analysis(self, doc, document_id=None):
        if document_id:
            labeled = list(self.db.labeled_data.find({"source_document_id": document_id}))
        else:
            labeled = list(self.db.labeled_data.find().limit(100))
        if not labeled:
            doc.add_paragraph("No transaction data available.")
            return
        df = pd.DataFrame(self._clean_numeric_data(labeled))
        if 'debit' in df.columns:
            doc.add_paragraph(f"Total Debit Amount: ${df['debit'].sum():,.2f}")
        if 'credit' in df.columns:
            doc.add_paragraph(f"Total Credit Amount: ${df['credit'].sum():,.2f}")
        # table
        doc.add_heading('Sample Transactions', level=2)
        table = doc.add_table(rows=1, cols=5)
        hdr = table.rows[0].cells
        headers = ['Date','Description','Debit','Credit','Category']
        for i,h in enumerate(headers):
            hdr[i].text = h
        for _, row in df.head(10).iterrows():
            r = table.add_row().cells
            r[0].text = str(row.get('date',''))
            r[1].text = str(row.get('description',''))[:50]
            d = row.get('debit',0)
            c = row.get('credit',0)
            r[2].text = f"${d:,.2f}" if d and d>0 else "-"
            r[3].text = f"${c:,.2f}" if c and c>0 else "-"
            r[4].text = str(row.get('category',''))

    def _add_categories_analysis(self, doc, document_id=None):
        if document_id:
            labeled = list(self.db.labeled_data.find({"source_document_id": document_id}))
        else:
            labeled = list(self.db.labeled_data.find())
        if not labeled:
            doc.add_paragraph("No category data available.")
            return
        df = pd.DataFrame(self._clean_numeric_data(labeled))
        if 'category' in df.columns:
            counts = df['category'].value_counts()
            doc.add_paragraph("Transaction Count by Category:")
            for cat, cnt in counts.items():
                doc.add_paragraph(f"- {cat}: {cnt}", style='List Bullet')
            if 'debit' in df.columns:
                doc.add_heading('Category Spending Analysis', level=2)
                spend = df.groupby('category')['debit'].sum().sort_values(ascending=False)
                for cat, amt in spend.items():
                    if amt>0:
                        doc.add_paragraph(f"- {cat}: ${amt:,.2f}", style='List Bullet')

    def _clean_numeric_data(self, labeled_data):
        cleaned = []
        for it in labeled_data:
            ci = {}
            for k,v in it.items():
                if k == '_id': continue
                if k in ['debit','credit','balance']:
                    try:
                        if v is None or v == '':
                            ci[k] = 0.0
                        elif isinstance(v, str):
                            s = v.replace('$','').replace(',','').strip()
                            ci[k] = float(s) if s else 0.0
                        else:
                            ci[k] = float(v)
                    except Exception:
                        ci[k] = 0.0
                else:
                    ci[k] = v
            cleaned.append(ci)
        return cleaned
